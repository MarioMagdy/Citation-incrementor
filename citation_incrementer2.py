from docx import Document
import re
from tkinter import Tk, filedialog
import time


# Define a function to increment citations
def fix_citations(match):
    citation_number = int(match.group()[1:-1])
    return f"[{citation_number + amount}]"

print("""Make sure you understand...
      this program is used to increament the citation meaning if you add more refereces to the start
      of the references list this will help you, you can count how many you will add and run this
      program on the latter one to add the amount to each number in each citation in it
      like if you will add 8 references you will need the referenece number 1 to be 9...
      -------------------------------------------------------------------------------------------------
      IMPORTANT
      this program finds only numbers between brackets like [3] or [21] or [2131]
      it won't work if [231 ] nor [1, 1] nor [123, 213] nor...

      it may remove the images -sorry- :(
      -------------------------------------------------------------------------------------------------
      """)


input("To Start press Enter...")
l=1
while l == 1 :


    # Create a Tkinter root window
    root = Tk()
    
    # Position the window at (x=600, y=20)
    root.geometry('+600+20')
    
    root.withdraw()  # Hide the root window
    

    root.focus_force()

    # Ask the user to select a file using a file dialog
    file_path = filedialog.askopenfilename(title="Select a Word document to increament", filetypes=[("Word files", "*.docx")])
    root.focus_force()
    root.focus_set()

    print("reading...")
    time.sleep(0.5)

    # Load the selected document
    doc = Document(file_path)

    # file = input("""File name?  \t Make sure it's in the same folder as the program and it's ".docx" \n""")
    # doc = Document(file+".docx")
    all_paras = doc.paragraphs
    text = ''
    x= 1
    
    while x== 1:
        try:
            amount = int(input("Amount to increament:\n"))
            x= 0
        except ValueError:
            print('Please enter an integer!')


    print("Working on it...")
    time.sleep(1.5)

    # Iterate through paragraphs and runs to modify text
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            run.text = re.sub(r"\[\d+\]", fix_citations, run.text)


    print("Saving...")
    time.sleep(0.5)

    output = ".".join([*file_path.split('.')[:-1]])+"_Increamented.docx"
    # Save the modified document
    doc.save(output)

    print(f"Saved to '{output}'")
    time.sleep(0.5)

    print("Done.\n\n")
    time.sleep(0.5)

    des =input("keep going? (y,n)")
    l = des.lower() == 'y'
